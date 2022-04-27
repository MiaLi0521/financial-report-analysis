import os

import docx
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


from .utils import read_company_code


class ReportDocument:
    def __init__(self, analysis, doc_name=None):
        self.doc_name = doc_name
        self.titles = analysis.titles
        self.tables = analysis.format_tables
        self.images = analysis.images
        self.table_index = analysis.table_index
        self.doc = docx.Document()

        self.doc.styles['Normal'].font.name = u'宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        self.doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        if self.doc_name is None:
            self.doc_name = \
                f"{read_company_code()} {analysis.config_file.split('.')[0].upper()} " \
                f"{self.tables['t1'].columns[1]}~{self.tables['t1'].columns[-1]}.docx"

    def _save_table(self, table):
        rows, cols = table.shape[0] + 1, table.shape[1]
        t = self.doc.add_table(rows, cols)

        # add the header rows.
        for col in range(cols - 1):
            t.cell(0, col + 1).text = str(table.columns[col + 1])

        # add the index cols.
        for row in range(rows - 1):
            t.cell(row + 1, 0).text = str(table.index[row])

        # add the rest of the dataFrame
        for row in range(rows - 1):
            for col in range(cols - 1):
                t.cell(row + 1, col + 1).text = str(table.values[row, col + 1])

        self.doc.add_paragraph()
        self.doc.save(os.path.join('dist', self.doc_name))

    def _save_title(self, title):
        head = self.doc.add_heading("", level=1)
        run_title = head.add_run(title)
        run_title.font.name = u'Cambria'
        run_title.font.size = Pt(18)
        run_title.font.color.rgb = RGBColor(31, 119, 180)
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
        self.doc.add_paragraph()
        self.doc.save(os.path.join('dist', self.doc_name))

    def _save_image(self, image):
        if isinstance(image, list):
            for img in image:
                self.doc.add_picture(os.path.join('images', f'{img}.png'))
        elif image:
            self.doc.add_picture(os.path.join('images', f'{image}.png'))
        self.doc.save(os.path.join('dist', self.doc_name))

    def save(self):
        for table_index in self.table_index:
            self._save_title(self.titles[table_index])
            self._save_table(self.tables[table_index])
            if self.images.get(table_index):
                self._save_image(self.images[table_index])

